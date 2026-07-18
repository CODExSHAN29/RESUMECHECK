"""One-off script to (re)generate sample_resumes/ fixtures used by tests
and manual verification. Run with: python tests/generate_fixtures.py
Requires requirements-dev.txt (adds reportlab, only needed for this script).
"""

import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

ROOT = os.path.join(os.path.dirname(__file__), "..", "..", "sample_resumes")
os.makedirs(ROOT, exist_ok=True)


def path(name: str) -> str:
    return os.path.join(ROOT, name)


# ---------- DOCX fixtures ----------

def build_clean_docx():
    doc = Document()
    doc.add_paragraph("Jordan Lee")
    doc.add_paragraph("jordan.lee@email.com | (555) 123-4567 | Austin, TX")

    doc.add_paragraph("Experience")
    doc.add_paragraph("Software Engineering Intern — Acme Corp", style=None)
    doc.add_paragraph("Jun 2024 - Aug 2024")
    doc.add_paragraph("• Built a data pipeline processing 1M records daily")
    doc.add_paragraph("• Reduced API latency by 30% through caching")
    doc.add_paragraph("• Collaborated with a team of 5 engineers using Agile")

    doc.add_paragraph("Education")
    doc.add_paragraph("B.S. Computer Science, University of Texas")
    doc.add_paragraph("Aug 2021 - May 2025")

    doc.add_paragraph("Skills")
    doc.add_paragraph("Python, JavaScript, React, SQL, Git, AWS")

    doc.add_paragraph("Projects")
    doc.add_paragraph("Resume Checker — Personal Project")
    doc.add_paragraph("• Built a full-stack web app with FastAPI and Next.js")
    doc.add_paragraph("• Implemented ATS-style parsing checks")

    doc.save(path("clean_one_column.docx"))


def build_multi_column_docx():
    doc = Document()
    doc.add_paragraph("Taylor Kim")
    doc.add_paragraph("taylor.kim@email.com | (555) 987-6543")

    section = doc.sections[0]
    sect_pr = section._sectPr
    existing_cols = sect_pr.find(qn("w:cols"))
    if existing_cols is not None:
        sect_pr.remove(existing_cols)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    sect_pr.append(cols)

    doc.add_paragraph("Experience")
    doc.add_paragraph("Marketing Intern — Beta Inc, Jan 2024 - May 2024")
    doc.add_paragraph("Education")
    doc.add_paragraph("B.A. Marketing, State University")
    doc.add_paragraph("Skills")
    doc.add_paragraph("SEO, Content Writing, Analytics")

    doc.save(path("multi_column.docx"))


def build_table_heavy_docx():
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Morgan Diaz"
    table.cell(0, 1).text = "morgan.diaz@email.com | (555) 222-3333"
    table.cell(1, 0).text = "Experience"
    table.cell(1, 1).text = "Sales Associate — Retail Co, 2022-2023"
    table.cell(2, 0).text = "Education"
    table.cell(2, 1).text = "A.A. Business, Community College"
    table.cell(3, 0).text = "Skills"
    table.cell(3, 1).text = "Excel, Customer Service, POS Systems"

    doc.save(path("table_heavy.docx"))


def build_header_contact_docx():
    doc = Document()
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "Casey Nguyen | casey.nguyen@email.com | (555) 444-5555"

    doc.add_paragraph("Experience")
    doc.add_paragraph("Barista — Coffee Shop, 2023-2024")
    doc.add_paragraph("Education")
    doc.add_paragraph("High School Diploma")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Customer Service, Cash Handling")

    doc.save(path("header_contact_only.docx"))


# ---------- PDF fixtures ----------

def build_clean_pdf():
    c = canvas.Canvas(path("clean_resume.pdf"), pagesize=LETTER)
    width, height = LETTER
    y = height - 100
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, y, "Alex Rivera")
    y -= 20
    c.setFont("Helvetica", 11)
    c.drawString(72, y, "alex.rivera@email.com | (555) 111-2222 | Denver, CO")
    y -= 30
    lines = [
        "Experience",
        "Data Analyst Intern - Gamma LLC",
        "Jun 2023 - Aug 2023",
        "- Analyzed sales data across 10 regions",
        "- Built dashboards used by 20+ stakeholders",
        "Education",
        "B.S. Statistics, Colorado State University",
        "Aug 2020 - May 2024",
        "Skills",
        "Python, R, SQL, Tableau, Excel",
    ]
    for line in lines:
        c.drawString(72, y, line)
        y -= 18
    c.save()


def build_multi_column_pdf():
    c = canvas.Canvas(path("multi_column.pdf"), pagesize=LETTER)
    width, height = LETTER
    y = height - 100
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, y, "Riley Chen")
    y -= 30
    c.setFont("Helvetica", 10)
    left_lines = [
        "Experience: Intern at Delta Co",
        "Responsibilities included data entry",
        "and client communication tasks daily",
        "Education: State University degree",
        "Bachelor of Arts in Communications",
    ]
    right_lines = [
        "Skills: Excel, PowerPoint, Outlook",
        "Languages: English, Spanish fluent",
        "Awards: Dean's List three semesters",
        "Volunteer: Local food bank helper",
        "Certifications: Google Analytics cert",
    ]
    for left, right in zip(left_lines, right_lines):
        c.drawString(72, y, left)
        c.drawString(320, y, right)
        y -= 20
    c.save()


def build_table_pdf():
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

    doc = SimpleDocTemplate(path("table_heavy.pdf"), pagesize=LETTER)
    data = [
        ["Sam Patel", "sam.patel@email.com"],
        ["Experience", "Intern - Epsilon Corp, 2023-2024"],
        ["Education", "B.S. Engineering, Tech University"],
        ["Skills", "AutoCAD, MATLAB, SolidWorks"],
    ]
    table = Table(data, colWidths=[150, 300])
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, (0, 0, 0))]))
    doc.build([table])


def build_scanned_image_pdf():
    c = canvas.Canvas(path("scanned_image_like.pdf"), pagesize=LETTER)
    width, height = LETTER
    # No text at all -- just vector shapes, simulating a scanned/flattened resume
    c.rect(50, 50, width - 100, height - 100)
    c.line(50, height / 2, width - 50, height / 2)
    c.showPage()
    c.save()


def build_header_contact_pdf():
    c = canvas.Canvas(path("header_contact_only.pdf"), pagesize=LETTER)
    width, height = LETTER
    c.setFont("Helvetica", 10)
    # Inside the top 8% margin (page is 792pt tall -> top 63pt)
    c.drawString(72, height - 30, "jamie.wong@email.com | (555) 666-7777")

    y = height - 120
    c.setFont("Helvetica-Bold", 12)
    c.drawString(72, y, "Jamie Wong")
    y -= 30
    for line in ["Experience", "Cashier - Grocery Store, 2022-2023", "Education", "High School Diploma", "Skills", "Teamwork, POS Systems"]:
        c.drawString(72, y, line)
        y -= 18
    c.save()


if __name__ == "__main__":
    build_clean_docx()
    build_multi_column_docx()
    build_table_heavy_docx()
    build_header_contact_docx()
    build_clean_pdf()
    build_multi_column_pdf()
    build_table_pdf()
    build_scanned_image_pdf()
    build_header_contact_pdf()
    print("Fixtures generated in", os.path.abspath(ROOT))
