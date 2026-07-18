import io

from docx import Document
from docx.oxml.ns import qn

from .common import ParsedDocument

WORDS_PER_PAGE_ESTIMATE = 500


def _iter_table_text(doc: Document) -> list[str]:
    lines = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        lines.append(p.text)
    return lines


def _collect_fonts(doc: Document) -> set[str]:
    fonts: set[str] = set()
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name:
                fonts.add(run.font.name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.name:
                            fonts.add(run.font.name)
    return fonts


def _section_column_count(section) -> int:
    sect_pr = section._sectPr
    cols_elements = sect_pr.findall(qn("w:cols"))
    if not cols_elements:
        return 1
    cols = cols_elements[-1]  # last one wins if a document has duplicates
    num = cols.get(qn("w:num"))
    try:
        return int(num) if num else 1
    except ValueError:
        return 1


def parse_docx(data: bytes, filename: str) -> ParsedDocument:
    doc = Document(io.BytesIO(data))

    body_lines = [p.text for p in doc.paragraphs if p.text.strip()]
    table_lines = _iter_table_text(doc)
    all_lines = body_lines + table_lines
    raw_text = "\n".join(all_lines)
    word_count = len(raw_text.split())
    char_count = len(raw_text)

    fonts = _collect_fonts(doc)

    table_count = len(doc.tables)

    multi_column_pages = 0
    for section in doc.sections:
        if _section_column_count(section) > 1:
            multi_column_pages += 1

    header_footer_lines: list[str] = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text.strip():
                header_footer_lines.append(p.text)
        for p in section.footer.paragraphs:
            if p.text.strip():
                header_footer_lines.append(p.text)

    image_count = len(doc.inline_shapes)
    text_box_count = len(list(doc.element.body.iter(qn("w:txbxContent"))))

    page_count = max(1, round(word_count / WORDS_PER_PAGE_ESTIMATE) or 1)

    return ParsedDocument(
        file_type="docx",
        filename=filename,
        raw_text=raw_text,
        lines=all_lines,
        body_text=raw_text,
        page_count=page_count,
        char_count=char_count,
        word_count=word_count,
        fonts=fonts,
        has_tables=table_count > 0,
        table_count=table_count,
        multi_column_pages=multi_column_pages,
        header_footer_text="\n".join(header_footer_lines),
        image_count=image_count,
        text_box_count=text_box_count,
        is_scanned_image=False,
    )
